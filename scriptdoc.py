"""The sales script as a formatted Word document.

Slack's code fence was the wrong container for this. It is the artefact a
closer reads immediately before a call -- often on a phone, often scrolling --
and monospace in a collapsed code block is the worst possible presentation for
that. A .docx opens in anything, prints, and can be annotated during the call.

Structure mirrors the order a closer actually needs it in: who they are, what
they said, what we found, what to ask, what to expect, how to present, and the
number. The price block is a table because it is the one thing that gets read
off the page verbatim under pressure.

Brand colours match the audit PDF so the two artefacts look like one pack.
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import pricing
import salesscript

PINE = RGBColor(0x1E, 0x3A, 0x2E)
BRASS = RGBColor(0xA9, 0x87, 0x4E)
INK = RGBColor(0x15, 0x14, 0x0F)
MUTED = RGBColor(0x6B, 0x65, 0x55)


def _style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)


def _heading(doc, text, size=13, color=PINE, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _kv(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    lab = p.add_run("%s  " % label)
    lab.bold = True
    lab.font.color.rgb = MUTED
    lab.font.size = Pt(9.5)
    p.add_run(str(value))
    return p


def _bullet(doc, text, bold_lead=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def _note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    return p


def _lead_field(lead, key, default=""):
    if lead is None:
        return default
    if isinstance(lead, dict):
        return lead.get(key) or default
    return getattr(lead, key, default) or default


def build(path, lead, *, evidence=None, dossier=None, recommendation=None,
          findings=(), business_name=""):
    """Write the sales script to `path` as .docx. Returns the path."""
    doc = Document()
    _style(doc)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.7)
        section.left_margin = section.right_margin = Inches(0.8)

    name = _lead_field(lead, "name", "(unknown)")
    company = business_name or _lead_field(lead, "domain", "")
    finding_keys = [f["key"] for f in findings]

    # --- title ---
    t = doc.add_paragraph()
    t.paragraph_format.space_after = Pt(0)
    run = t.add_run("Sales Call Brief")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = PINE
    sub = doc.add_paragraph()
    r = sub.add_run("%s  ·  %s" % (name, company))
    r.font.size = Pt(11)
    r.font.color.rgb = BRASS
    _note(doc, "Internal. Every figure is measured or quoted; anything not "
               "established is marked unknown. Do not fill a gap on the call.")

    # --- who ---
    _heading(doc, "Who you are speaking to")
    _kv(doc, "Contact", name)
    web = (dossier or {}).get("web_research") or {}
    role = (web.get("role") or {})
    if role.get("value"):
        _kv(doc, "Role", role["value"])
    _kv(doc, "Company", company)
    for label, key in (("Phone", "phone"), ("Call time", "appointment_at"),
                       ("Closer", "manager")):
        v = _lead_field(lead, key)
        if v:
            _kv(doc, label, v)
    _kv(doc, "Track", (_lead_field(lead, "track") or "local").upper())

    # --- background ---
    company_facts = (dossier or {}).get("company") or {}
    known = [(k, v) for k, v in company_facts.items()
             if isinstance(v, dict) and v.get("value") is not None]
    if known or dossier:
        _heading(doc, "Company background")
        for key, field in known:
            _kv(doc, key.replace("_", " ").title(),
                salesscript._readable(field.get("value")))
        unknown = (dossier or {}).get("unknown_fields") or []
        if unknown:
            _note(doc, "Not establishable from outside, confirm on the call: "
                       + ", ".join(u.replace("_", " ") for u in unknown))

    # --- what they told us ---
    _heading(doc, "What they told us")
    for label, key in (("Frustration", "frustration"),
                       ("Already tried", "tried"),
                       ("Budget answer", "budget"),
                       ("Revenue", "revenue"),
                       ("Decision process", "decision_role"),
                       ("Urgency", "urgency")):
        v = _lead_field(lead, key)
        _kv(doc, label, v if v else "(not answered)")

    # --- findings ---
    _heading(doc, "What the audit found")
    if findings:
        for f in findings:
            _bullet(doc, f["headline"])
    else:
        _bullet(doc, "No finding cleared its threshold. Run this as a "
                     "discovery call and do not overclaim.")

    # --- discovery ---
    _heading(doc, "Discovery questions")
    for q in salesscript.discovery_for(finding_keys,
                                       _lead_field(lead, "frustration")):
        _bullet(doc, q)

    # --- objections ---
    _heading(doc, "Objection handling")
    got = salesscript.objections_for(_lead_field(lead, "tried"))
    if got:
        for label, handling in got:
            _bullet(doc, handling, bold_lead="%s — " % label)
    else:
        _bullet(doc, "Nothing flagged in their 'already tried' answer.")

    # --- impact ---
    _heading(doc, "Expected business impact")
    anchor = (recommendation or {}).get("anchor_price")
    for line in salesscript.impact_lines(evidence, finding_keys, anchor=anchor):
        _bullet(doc, line)

    # --- how to present ---
    _heading(doc, "How to present the offer")
    for i, step in enumerate(salesscript.PRESENT_STEPS, 1):
        _bullet(doc, step, bold_lead="%d. " % i)

    # --- price ---
    if recommendation:
        _heading(doc, "Price")
        rec = recommendation
        cur = rec["currency"]
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0].cells
        for i, text in enumerate(("Tier", "Per month", "3 months upfront",
                                  "")):
            hdr[i].text = ""
            r = hdr[i].paragraphs[0].add_run(text)
            r.bold = True
            r.font.size = Pt(9.5)
        for tier in pricing.TIERS:
            row = table.add_row().cells
            row[0].text = tier.title()
            row[1].text = "%s %s" % (cur, "{:,}".format(rec["prices"][tier]))
            row[2].text = "%s %s" % (cur, "{:,}".format(rec["upfront"][tier]))
            marker = ("ANCHOR — open here" if tier == rec["anchor_tier"]
                      else ("step down" if tier == "foundation" else "step up"))
            cell = row[3]
            cell.text = ""
            r = cell.paragraphs[0].add_run(marker)
            r.font.size = Pt(9)
            if tier == rec["anchor_tier"]:
                r.bold = True
                r.font.color.rgb = BRASS

        doc.add_paragraph()
        _kv(doc, "Size class", "%s — %s" % (rec["size_class"],
                                            rec["size_basis"]))
        _kv(doc, "Push to Dominate", rec["push_to_dominate"])
        _kv(doc, "Upfront terms", rec["upfront_terms"])
        for flag in rec["flags"]:
            _note(doc, "! " + flag)
        if rec["unknown_signals"]:
            _note(doc, "Size signals not established: "
                       + ", ".join(rec["unknown_signals"]))

    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    doc.save(path)
    return path
